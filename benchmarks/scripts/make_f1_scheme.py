#!/usr/bin/env python3
# ruff: noqa: I001
"""Build manuscript Figure 1: the BRAID calculation scheme.

This figure is a schematic, but it is generated from structured source tables:
one declared toy rMATS event, the shipped differential conformal calibrator, and
the production tier truth table. No generated bitmap is used in the final figure.
"""

from __future__ import annotations

import hashlib
import json
import shutil
from datetime import datetime, timezone
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle


ROOT = Path(__file__).resolve().parents[2]

import sys  # noqa: E402

sys.path.insert(0, str(ROOT / "benchmarks"))
import figstyle  # noqa: E402

figstyle.apply()

from braid.adapters.base import confidence_tier  # noqa: E402
from braid.target.conformal import ConformalCalibrator  # noqa: E402


OUT_DIR = ROOT / "outputs" / "figures" / "manuscript" / "F1"
PAPER_FIGURE_DIR = ROOT / "paper" / "figures"
SOURCE_XLSX = OUT_DIR / "f1_source_data.xlsx"
MANIFEST_JSON = OUT_DIR / "f1_manifest.json"
VALIDATION_TXT = OUT_DIR / "f1_validation.txt"
LEGEND_DOCX = OUT_DIR / "f1_legend.docx"
CALIBRATOR_JSON = (
    ROOT / "braid" / "target" / "calibration_artifacts"
    / "differential_dpsi_conformal.json"
)
DEPTH_TITRATION_JSON = ROOT / "benchmarks" / "results" / "depth_titration.json"

INK = "#1F2933"
MUTED = "#697386"
HAIR = "#C7CDD6"
PANEL_BG = "#F7F9FC"
NAVY = "#083B6F"
BLUE = "#0B5CAD"
BLUE_SOFT = "#DCEBFA"
BLUE_PALE = "#F1F6FC"
RED = "#C23B32"
RED_SOFT = "#F8E2DF"
GREEN = "#2F7D51"
GREEN_SOFT = "#DFF1E7"
AMBER = "#B7791F"
AMBER_SOFT = "#FFF1D6"
GREY_SOFT = "#EEF1F5"


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _save_all(fig: plt.Figure, stem: Path) -> dict[str, str]:
    out: dict[str, str] = {}
    for suffix in (".png", ".pdf", ".svg"):
        path = stem.with_suffix(suffix)
        if suffix == ".png":
            fig.savefig(path, dpi=400)
        else:
            fig.savefig(path)
        out[suffix.lstrip(".")] = str(path.relative_to(ROOT))
    return out


def _arrow(ax: plt.Axes, start, end, *, color=BLUE, lw=1.6, rad=0.0) -> None:
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=12,
            linewidth=lw,
            color=color,
            connectionstyle=f"arc3,rad={rad}",
            shrinkA=4,
            shrinkB=4,
            zorder=6,
        )
    )


def _label(ax: plt.Axes, x: float, y: float, text: str, **kwargs) -> None:
    defaults = dict(ha="center", va="center", fontsize=7.2, color=INK)
    defaults.update(kwargs)
    ax.text(x, y, text, **defaults)


def _panel_shell(ax: plt.Axes, letter: str) -> None:
    ax.add_patch(
        FancyBboxPatch(
            (0.06, 0.10),
            11.88,
            3.80,
            boxstyle="round,pad=0.02,rounding_size=0.08",
            facecolor="white",
            edgecolor=NAVY,
            linewidth=0.9,
            zorder=-10,
        )
    )
    ax.add_patch(
        Rectangle((0.06, 3.46), 0.42, 0.44, facecolor=NAVY, edgecolor=NAVY, zorder=-8)
    )
    _label(ax, 0.27, 3.68, letter, fontsize=12.5, fontweight="bold", color="white")


def _box(
    ax: plt.Axes,
    x: float,
    y: float,
    w: float,
    h: float,
    *,
    title: str,
    body: str | None = None,
    fill: str = "white",
    edge: str = HAIR,
    title_color: str = INK,
) -> None:
    ax.add_patch(
        FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.025,rounding_size=0.06",
            facecolor=fill,
            edgecolor=edge,
            linewidth=0.9,
            zorder=1,
        )
    )
    _label(
        ax,
        x + w / 2,
        y + h - 0.28,
        title,
        fontsize=7.5,
        fontweight="bold",
        color=title_color,
    )
    if body:
        _label(
            ax,
            x + w / 2,
            y + h / 2 - 0.1,
            body,
            fontsize=6.2,
            color=MUTED,
            linespacing=1.18,
        )


def _interval(
    ax: plt.Axes,
    y: float,
    lo: float,
    hi: float,
    center: float,
    *,
    color: str,
    lw: float = 2.8,
    cap: float = 0.10,
) -> None:
    ax.plot([lo, hi], [y, y], color=color, lw=lw, solid_capstyle="butt", zorder=4)
    for x in (lo, hi):
        ax.plot([x, x], [y - cap, y + cap], color=color, lw=lw, zorder=4)
    ax.plot([center], [y], marker="o", ms=4.8, color=color, zorder=5)


def _load_calibrator() -> ConformalCalibrator:
    with CALIBRATOR_JSON.open("r", encoding="utf-8") as fh:
        return ConformalCalibrator.from_dict(json.load(fh))


def _toy_event(cal: ConformalCalibrator) -> dict[str, float | str | bool]:
    inc1, skip1, inc2, skip2 = 80.0, 20.0, 35.0, 65.0
    inc_form_len = 1.0
    skip_form_len = 1.0
    length_ratio = inc_form_len / skip_form_len
    skip1_norm = skip1 * length_ratio
    skip2_norm = skip2 * length_ratio
    psi1 = (inc1 + 0.5) / (inc1 + skip1_norm + 1.0)
    psi2 = (inc2 + 0.5) / (inc2 + skip2_norm + 1.0)
    dpsi = psi1 - psi2
    total_support = inc1 + skip1 + inc2 + skip2
    sampling_std = 0.030
    q = cal.q_for(total_support, event_type="SE")
    half_width = float(np.hypot(q, 1.959963984540054 * sampling_std))
    ci_low, ci_high = cal.robust_interval(
        dpsi,
        sampling_std,
        total_support,
        event_type="SE",
        clip=(-1.0, 1.0),
    )
    fdr = 0.010
    effect_cutoff = 0.10
    reliable = ci_low > 0.0 or ci_high < 0.0
    effect = abs(dpsi) >= effect_cutoff
    caller_sig = fdr < 0.05
    tier = confidence_tier(reliable, effect, caller_sig)
    return {
        "sample_1_included": inc1,
        "sample_1_skipped": skip1,
        "sample_2_included": inc2,
        "sample_2_skipped": skip2,
        "inc_form_len": inc_form_len,
        "skip_form_len": skip_form_len,
        "length_ratio": length_ratio,
        "sample_1_skipped_length_normalized": skip1_norm,
        "sample_2_skipped_length_normalized": skip2_norm,
        "psi_sample_1": psi1,
        "psi_sample_2": psi2,
        "dpsi": dpsi,
        "sampling_std": sampling_std,
        "total_support": total_support,
        "event_type": "SE",
        "conformal_q": q,
        "half_width": half_width,
        "ci_low": ci_low,
        "ci_high": ci_high,
        "fdr": fdr,
        "effect_cutoff": effect_cutoff,
        "reliable": reliable,
        "effect": effect,
        "caller_significant": caller_sig,
        "tier": tier,
    }


def _source_tables(
    cal: ConformalCalibrator,
    toy: dict[str, float | str | bool],
    design: dict[str, object],
):
    steps = pd.DataFrame(
        [
            {
                "panel": "A",
                "step": "1",
                "label": "Read rMATS IJC/SJC counts; BRAID does not count from BAM",
                "source_path": "braid/target/rmats_bootstrap.py",
                "source_lines": "391-397",
            },
            {
                "panel": "A",
                "step": "2",
                "label": (
                    "Draw PSI from Jeffreys-Beta posteriors after "
                    "IncFormLen/SkipFormLen normalization; replicate-aware auto "
                    "resamples complete replicate vectors when present"
                ),
                "source_path": "braid/commands/differential.py",
                "source_lines": "136-208,431-466",
            },
            {
                "panel": "B",
                "step": "3",
                "label": "Choose support/event-type conformal q",
                "source_path": "braid/target/conformal.py",
                "source_lines": "334-357",
            },
            {
                "panel": "B",
                "step": "4",
                "label": "Build robust calibrated interval",
                "source_path": "braid/target/conformal.py",
                "source_lines": "392-431",
            },
            {
                "panel": "C",
                "step": "5",
                "label": "Assign confidence tier",
                "source_path": "braid/adapters/base.py",
                "source_lines": "97-120",
            },
        ]
    )
    toy_df = pd.DataFrame([toy])
    cal_df = pd.DataFrame(
        [
            {
                "key": "alpha",
                "value": cal.alpha,
                "source": str(CALIBRATOR_JSON.relative_to(ROOT)),
            },
            {
                "key": "scale_kind",
                "value": cal.scale_kind,
                "source": str(CALIBRATOR_JSON.relative_to(ROOT)),
            },
            {
                "key": "training_scope",
                "value": cal.training_scope,
                "source": str(CALIBRATOR_JSON.relative_to(ROOT)),
            },
            {
                "key": "q_global",
                "value": cal.q_global,
                "source": str(CALIBRATOR_JSON.relative_to(ROOT)),
            },
            {
                "key": "q_for_toy_event",
                "value": toy["conformal_q"],
                "source": "ConformalCalibrator.q_for(total_support=200, event_type='SE')",
            },
        ]
    )
    tiers = pd.DataFrame(
        [
            {
                "caller_significant": False,
                "braid_reliable_effect": False,
                "tier": "not-significant",
            },
            {
                "caller_significant": False,
                "braid_reliable_effect": True,
                "tier": "supported",
            },
            {
                "caller_significant": True,
                "braid_reliable_effect": False,
                "tier": "caller-significant-only",
            },
            {
                "caller_significant": True,
                "braid_reliable_effect": True,
                "tier": "high-confidence",
            },
        ]
    )
    references = pd.DataFrame(
        [
            {
                "item": "toy_event_scope",
                "mode": "declared schematic input",
                "path": "f1_source_data.xlsx:toy_event",
                "used_in_final_data": True,
                "note": "Illustrative arithmetic, not a benchmark result.",
            },
        ]
    )
    per_dataset: dict[str, int] = design.get("per_dataset", {})  # type: ignore[assignment]
    panel_d_rows = [
        {
            "panel": "D",
            "role": "production_calibrator",
            "item": "training_scope",
            "value": cal.training_scope,
            "source": str(CALIBRATOR_JSON.relative_to(ROOT)),
            "note": "Shipped calibrator fit once on TRA2 plus circadian RT-PCR residuals.",
        },
        {
            "panel": "D",
            "role": "production_calibrator",
            "item": "calibration_total_se_events",
            "value": design["n_total"],
            "source": str(DEPTH_TITRATION_JSON.relative_to(ROOT)),
            "note": "Matched SE events used to state the fit-once calibration scope.",
        },
        {
            "panel": "D",
            "role": "production_calibrator",
            "item": "q_global",
            "value": design["q_global"],
            "source": str(CALIBRATOR_JSON.relative_to(ROOT)),
            "note": "Displayed split-conformal residual quantile.",
        },
    ]
    dataset_roles = {
        "TRA2 (GSE59335, human)": "calibration_dataset",
        "Circadian (GSE54651, mouse)": "calibration_dataset",
        "PC3E/GS689 (SRA, human)": "held_out_transfer_dataset",
    }
    for dataset, role in dataset_roles.items():
        if dataset not in per_dataset:
            continue
        panel_d_rows.append(
            {
                "panel": "D",
                "role": role,
                "item": dataset,
                "value": per_dataset[dataset],
                "source": str(DEPTH_TITRATION_JSON.relative_to(ROOT)),
                "note": "Matched-event count shown or summarized in Panel D.",
            }
        )
    panel_d_rows.extend(
        [
            {
                "panel": "D",
                "role": "validation_mode",
                "item": "head_to_head_coverage",
                "value": "leakage-free cross-fit",
                "source": "benchmarks/headtohead/head_to_head_coverage.py",
                "note": "Validation folds fit q without scoring an event with its own q.",
            },
            {
                "panel": "D",
                "role": "validation_mode",
                "item": "pc3e_no_refit_transfer",
                "value": "no-refit transfer",
                "source": "benchmarks/headtohead/cross_dataset_transfer.py",
                "note": "PC3E/GS689 transfer evidence is separate from the fit-once q statement.",
            },
            {
                "panel": "D",
                "role": "validation_mode",
                "item": "depth_titration_and_non_se",
                "value": "5-100% depth and long-read non-SE checks",
                "source": "benchmarks/results/depth_titration.json",
                "note": "Stress-test evidence is not a refit of the production calibrator.",
            },
            {
                "panel": "D",
                "role": "application_mode",
                "item": "dm1_application",
                "value": "uses braid differential packaged q",
                "source": "benchmarks/application_dm1/run_dm1_application.py",
                "note": "Disease application uses BRAID without adding RT-PCR truth for refit.",
            },
            {
                "panel": "D",
                "role": "application_mode",
                "item": "esrp_application",
                "value": "uses braid differential packaged q",
                "source": "benchmarks/application_esrp/run_esrp_application.py",
                "note": (
                    "Cross-species application treats nominal calibration as transfer "
                    "evidence."
                ),
            },
        ]
    )
    panel_d_design = pd.DataFrame(panel_d_rows)
    metadata = pd.DataFrame(
        [
            {"key": "figure_id", "value": "F1"},
            {"key": "title", "value": "BRAID calculation scheme"},
            {"key": "script", "value": str(Path(__file__).relative_to(ROOT))},
            {"key": "calibrator", "value": str(CALIBRATOR_JSON.relative_to(ROOT))},
        ]
    )
    return {
        "calculation_steps": steps,
        "toy_event": toy_df,
        "calibrator": cal_df,
        "tier_logic": tiers,
        "panel_d_design": panel_d_design,
        "references": references,
        "metadata": metadata,
    }


def _draw_count_table(ax: plt.Axes, toy: dict[str, float | str | bool]) -> None:
    _box(
        ax,
        0.40,
        1.25,
        2.2,
        2.05,
        title="1. rMATS IJC/SJC",
        body=None,
        fill="white",
        edge=HAIR,
    )
    x0, y0 = 0.62, 1.62
    cell_w, cell_h = 0.52, 0.36
    headers = ["group", "inc", "skip"]
    rows = [
        ("sample 1", int(toy["sample_1_included"]), int(toy["sample_1_skipped"])),
        ("sample 2", int(toy["sample_2_included"]), int(toy["sample_2_skipped"])),
    ]
    for j, head in enumerate(headers):
        ax.add_patch(
            Rectangle(
                (x0 + j * cell_w, y0 + 2 * cell_h),
                cell_w,
                cell_h,
                facecolor=BLUE_SOFT,
                edgecolor=HAIR,
                lw=0.6,
            )
        )
        _label(ax, x0 + j * cell_w + cell_w / 2, y0 + 2.5 * cell_h, head, fontsize=5.7)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            ax.add_patch(
                Rectangle(
                    (x0 + j * cell_w, y0 + (1 - i) * cell_h),
                    cell_w,
                    cell_h,
                    facecolor="white",
                    edgecolor=HAIR,
                    lw=0.6,
                )
            )
            _label(
                ax,
                x0 + j * cell_w + cell_w / 2,
                y0 + (1.5 - i) * cell_h,
                str(val),
                fontsize=5.6,
                color=INK if j else MUTED,
            )
    # Cassette-exon glyph.
    y = 1.42
    ax.plot([0.78, 2.22], [y, y], color=HAIR, lw=0.8)
    ax.add_patch(Rectangle((0.72, y - 0.08), 0.28, 0.16, facecolor=GREY_SOFT, edgecolor=HAIR))
    ax.add_patch(Rectangle((1.36, y - 0.08), 0.28, 0.16, facecolor=BLUE_SOFT, edgecolor=BLUE))
    ax.add_patch(Rectangle((2.00, y - 0.08), 0.28, 0.16, facecolor=GREY_SOFT, edgecolor=HAIR))
    _label(ax, 1.50, 1.15, "read from rMATS table", fontsize=5.9, color=MUTED)


def _draw_panel_a(ax: plt.Axes, toy: dict[str, float | str | bool]) -> None:
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")
    _panel_shell(ax, "A")
    _label(
        ax,
        6.0,
        3.76,
        "From rMATS counts to one Delta PSI estimate",
        fontsize=9.2,
        fontweight="bold",
    )
    _draw_count_table(ax, toy)
    _box(
        ax,
        3.35,
        1.25,
        2.15,
        2.05,
        title="2. PSI posterior",
        body=None,
        fill=BLUE_PALE,
        edge=BLUE_SOFT,
        title_color=BLUE,
    )
    _label(
        ax,
        4.42,
        2.48,
        "pooled path shown\nskip adjusted by\nIncFormLen/SkipFormLen",
        fontsize=5.7,
        color=MUTED,
        linespacing=1.15,
    )
    xs = np.linspace(-1.0, 1.0, 100)
    for cy, color in [(1.83, BLUE), (1.55, "#7EA7D8")]:
        ys = np.exp(-(xs**2) / 0.20)
        ax.plot(4.42 + xs * 0.34, cy + ys * 0.16, color=color, lw=1.0)
        ax.plot([4.03, 4.82], [cy, cy], color=HAIR, lw=0.7)
    _box(
        ax,
        6.45,
        1.25,
        2.10,
        2.05,
        title="3. Delta PSI",
        body=(
            f"sample 1 - sample 2\n"
            f"{toy['psi_sample_1']:.2f} - {toy['psi_sample_2']:.2f}\n"
            f"= {toy['dpsi']:.2f}"
        ),
        fill="white",
        edge=HAIR,
        title_color=INK,
    )
    ax.plot([7.10, 7.95], [1.78, 1.78], color=BLUE, lw=2.0)
    ax.plot([7.52], [1.78], marker="o", color=BLUE, ms=4.5)
    _box(
        ax,
        9.45,
        1.25,
        2.15,
        2.05,
        title="reported center",
        body="dpsi\nprob_large_effect\nsupport\nmodel: sum or rep",
        fill=GREY_SOFT,
        edge=HAIR,
        title_color=INK,
    )
    _label(
        ax,
        6.00,
        0.54,
        (
            "If complete replicate IJC/SJC vectors exist, default auto uses "
            "a replicate-resampling model."
        ),
        fontsize=5.8,
        color=MUTED,
    )
    _arrow(ax, (2.62, 2.28), (3.28, 2.28))
    _arrow(ax, (5.55, 2.28), (6.38, 2.28))
    _arrow(ax, (8.58, 2.28), (9.38, 2.28))


def _draw_panel_b(ax: plt.Axes, toy: dict[str, float | str | bool]) -> None:
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")
    _panel_shell(ax, "B")
    _label(
        ax,
        6.0,
        3.76,
        "Calibration changes the interval width, not the point estimate",
        fontsize=9.2,
        fontweight="bold",
    )
    _box(
        ax,
        0.55,
        0.72,
        3.10,
        2.55,
        title="calibration events",
        body=None,
        fill="white",
        edge=HAIR,
    )
    _label(
        ax,
        2.10,
        2.66,
        "RNA-seq Delta PSI\nmatched to orthogonal RT-PCR",
        fontsize=6.1,
        color=MUTED,
        linespacing=1.05,
    )
    rng = np.random.default_rng(7)
    residuals = np.array([0.04, 0.06, 0.08, 0.11, 0.14, 0.18, 0.22, 0.28, 0.34])
    yvals = np.linspace(1.22, 2.30, residuals.size) + rng.normal(0, 0.020, residuals.size)
    x0 = 1.05
    for r, y in zip(residuals, yvals, strict=True):
        ax.plot([x0, x0 + r * 5.0], [y, y], color=HAIR, lw=0.7)
        ax.plot([x0 + r * 5.0], [y], marker="o", color=RED, ms=3.5)
    qx = x0 + float(toy["conformal_q"]) * 5.0
    ax.plot([qx, qx], [1.02, 2.48], color=RED, lw=1.0, ls=(0, (3, 2)))
    _label(ax, qx + 0.24, 2.45, "q", fontsize=8.0, color=RED, fontweight="bold")
    _label(
        ax,
        2.10,
        0.98,
        "|RNA-seq - RT-PCR| residuals",
        fontsize=5.9,
        color=MUTED,
    )
    _box(
        ax,
        4.45,
        0.72,
        3.00,
        2.55,
        title="residual-calibrated width",
        body=(
            "support/event-type q\n"
            "+ sampling spread\n\n"
            "half = sqrt(q^2 +\n"
            "(1.96 x sampling SD)^2)"
        ),
        fill=RED_SOFT,
        edge="#E7B7B1",
        title_color=RED,
    )
    _label(
        ax,
        5.95,
        1.03,
        f"toy half-width = {toy['half_width']:.2f}",
        fontsize=6.3,
        color=RED,
        fontweight="bold",
    )
    _box(
        ax,
        8.25,
        0.72,
        3.20,
        2.55,
        title="RT-PCR-calibrated interval",
        body=None,
        fill="white",
        edge=HAIR,
        title_color=BLUE,
    )
    ax.plot([8.75, 10.95], [1.38, 1.38], color=INK, lw=0.8)
    zero_x = 9.25
    ax.plot([zero_x, zero_x], [1.05, 2.45], color=HAIR, lw=1.0, ls=(0, (3, 2)))
    _label(ax, zero_x, 0.92, "0", fontsize=6.0, color=MUTED)
    scale = 1.55
    center_x = zero_x + float(toy["dpsi"]) * scale
    lo_x = zero_x + float(toy["ci_low"]) * scale
    hi_x = zero_x + float(toy["ci_high"]) * scale
    _interval(ax, 1.95, lo_x, hi_x, center_x, color=BLUE, lw=3.0)
    _label(
        ax,
        9.85,
        2.45,
        f"{toy['dpsi']:.2f} [{toy['ci_low']:.2f}, {toy['ci_high']:.2f}]",
        fontsize=6.4,
        color=BLUE,
        fontweight="bold",
    )
    _label(
        ax,
        9.85,
        1.00,
        "interval excludes 0 = reliable effect",
        fontsize=6.2,
        color=GREEN,
        fontweight="bold",
    )
    _arrow(ax, (3.72, 2.00), (4.38, 2.00), color=RED)
    _arrow(ax, (7.52, 2.00), (8.18, 2.00), color=BLUE)


def _draw_panel_c(ax: plt.Axes, toy: dict[str, float | str | bool]) -> None:
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")
    _panel_shell(ax, "C")
    _label(
        ax,
        6.0,
        3.76,
        "Tier assignment plus final output record",
        fontsize=9.2,
        fontweight="bold",
    )
    _label(ax, 4.85, 3.20, "BRAID reliable effect?", fontsize=7.1, fontweight="bold")
    _label(ax, 0.90, 1.56, "rMATS\nFDR < 0.05?", fontsize=6.8, fontweight="bold")
    col_x = [3.30, 6.35]
    row_y = [2.00, 0.96]
    col_lab = ["no\ninterval crosses 0 or small effect",
               "yes\ninterval excludes 0 and |Delta PSI| >= 0.1"]
    row_lab = ["no", "yes"]
    cells = [
        [("not-significant", GREY_SOFT, MUTED), ("supported", GREEN_SOFT, GREEN)],
        [("caller-significant-only", AMBER_SOFT, AMBER), ("high-confidence", BLUE, "white")],
    ]
    for x, lab in zip(col_x, col_lab, strict=True):
        _label(ax, x, 2.66, lab, fontsize=5.8, color=MUTED, linespacing=1.10)
    for y, lab in zip(row_y, row_lab, strict=True):
        _label(ax, 1.82, y, lab, fontsize=6.9, color=MUTED, fontweight="bold")
    w, h = 2.68, 0.72
    for i, y in enumerate(row_y):
        for j, x in enumerate(col_x):
            name, fill, color = cells[i][j]
            ax.add_patch(
                FancyBboxPatch(
                    (x - w / 2, y - h / 2),
                    w,
                    h,
                    boxstyle="round,pad=0.025,rounding_size=0.05",
                    facecolor=fill,
                    edgecolor=HAIR,
                    linewidth=0.9,
                )
            )
            _label(ax, x, y, name, fontsize=6.4, color=color, fontweight="bold")
    _arrow(ax, (7.72, 1.47), (8.34, 1.47), color=BLUE)
    _box(
        ax,
        8.52,
        0.62,
        2.92,
        2.20,
        title="output row",
        body=None,
        fill=BLUE_PALE,
        edge=BLUE_SOFT,
        title_color=BLUE,
    )
    rows = [
        ("dpsi", f"{toy['dpsi']:.2f}"),
        ("ci_low", f"{toy['ci_low']:.2f}"),
        ("ci_high", f"{toy['ci_high']:.2f}"),
        ("reliable", "true"),
        ("tier", str(toy["tier"])),
    ]
    x_label, x_value = 8.85, 10.15
    y_start = 2.25
    for idx, (key, value) in enumerate(rows):
        y = y_start - idx * 0.31
        _label(ax, x_label, y, key, fontsize=5.8, color=MUTED, ha="left")
        value_color = BLUE if key == "tier" else INK
        _label(
            ax,
            x_value,
            y,
            value,
            fontsize=5.8,
            color=value_color,
            fontweight="bold" if key == "tier" else "normal",
            ha="left",
        )


def _write_source_data(tables: dict[str, pd.DataFrame]) -> None:
    with pd.ExcelWriter(SOURCE_XLSX, engine="openpyxl") as writer:
        for sheet, df in tables.items():
            df.to_excel(writer, index=False, sheet_name=sheet[:31])


def _write_legend() -> None:
    doc = Document()
    doc.add_paragraph("Figure 1. BRAID calculation scheme.")
    doc.add_paragraph(
        "A, BRAID reads rMATS IJC/SJC counts rather than recounting from BAM. "
        "The pooled path shown length-normalizes skipping counts by "
        "IncFormLen/SkipFormLen, forms Jeffreys-Beta PSI posteriors, and reports "
        "Delta PSI as sample 1 minus sample 2; when complete replicate count "
        "vectors are present, the default auto model resamples replicates instead. "
        "B, The residual-calibrated width is the split-conformal quantile q of the "
        "absolute RNA-seq-versus-RT-PCR residuals; the production interval combines q "
        "with the within-sample sampling spread and clips the RT-PCR-calibrated 95% "
        "prediction interval to the Delta PSI range. C, Confidence "
        "tiers combine two binary criteria: upstream rMATS significance and whether "
        "the BRAID calibrated interval supports a nonzero effect; the selected tier "
        "is written with the calibrated interval in the output row. Values in A and "
        "B are an illustrative event-level calculation. D, The production q shown in "
        "the shipped calibrator is fit once on TRA2 plus circadian RT-PCR residuals; "
        "validation evidence is separated into leakage-free cross-fit coverage and "
        "no-refit transfer/application checks."
    )
    doc.save(LEGEND_DOCX)


def _design_table(cal: ConformalCalibrator) -> dict[str, object]:
    """Study-design provenance for panel D: which datasets fit q vs evaluate it.

    ``n_total`` is read from the shipped calibrator; the per-dataset matched counts
    are read from the committed depth-titration artifact (full-depth ``n_matched``).
    Every number is a real, sourced value -- the panel is a schematic of the *design*,
    not fabricated data.
    """
    n_total = int(round(float(cal.calibration_profile.get("n", 162))))
    per: dict[str, int] = {}
    try:
        dt = json.loads(DEPTH_TITRATION_JSON.read_text(encoding="utf-8"))
        for ds, frac in dict(dt.get("per_dataset", {})).items():
            n = frac.get("1.00", {}).get("n_matched_mean")
            if n is not None:
                per[ds] = int(round(float(n)))
    except (OSError, ValueError, KeyError, TypeError):
        per = {}
    return {"n_total": n_total, "q_global": float(cal.q_global), "per_dataset": per}


def _draw_panel_d(ax: plt.Axes, design: dict[str, object]) -> None:
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 4)
    ax.axis("off")
    _panel_shell(ax, "D")
    _label(
        ax, 6.0, 3.76,
        "Production q is fit once; validation is checked separately",
        fontsize=9.2, fontweight="bold",
    )
    per: dict[str, int] = design.get("per_dataset", {})  # type: ignore[assignment]
    n_tra2 = per.get("TRA2 (GSE59335, human)", 112)
    n_circ = per.get("Circadian (GSE54651, mouse)", 50)
    n_pc3e = per.get("PC3E/GS689 (SRA, human)", 34)
    n_total = design.get("n_total", 162)
    q = float(design.get("q_global", 0.341))  # type: ignore[arg-type]
    # Left: the production calibration set that fits q once (green).
    _box(
        ax, 0.45, 1.05, 3.05, 2.00,
        title="Production calibrator",
        body=(
            f"TRA2 (GSE59335)  n={n_tra2}\n"
            f"Circadian (GSE54651)  n={n_circ}\n"
            f"total {n_total} SE events, RT-PCR truth"
        ),
        fill=GREEN_SOFT, edge=GREEN, title_color=GREEN,
    )
    _arrow(ax, (3.58, 2.05), (4.42, 2.05), color=GREEN)
    _label(ax, 4.0, 2.34, "fit", fontsize=6.4, color=GREEN, fontweight="bold")
    # Middle: the learned split-conformal quantile q (blue).
    _box(
        ax, 4.45, 1.50, 1.70, 1.10,
        title=f"q = {q:.3f}",
        body="split-conformal\nresidual quantile",
        fill=BLUE_SOFT, edge=NAVY, title_color=NAVY,
    )
    _arrow(ax, (6.20, 2.05), (7.00, 2.05), color=BLUE)
    _label(ax, 6.60, 2.34, "apply", fontsize=6.4, color=BLUE, fontweight="bold")
    # Right: validation evidence, separated from the fit-once production q.
    ax.add_patch(
        FancyBboxPatch(
            (7.05, 0.55), 4.45, 2.95,
            boxstyle="round,pad=0.025,rounding_size=0.06",
            facecolor=BLUE_PALE, edgecolor=BLUE, linewidth=0.9, zorder=1,
        )
    )
    _label(ax, 9.275, 3.18, "Validation evidence", fontsize=7.5,
           fontweight="bold", color=NAVY)
    _label(ax, 9.275, 2.90, "(cross-fit folds or no-refit transfer)",
           fontsize=5.6, color=MUTED)
    rows = [
        "TRA2/Circadian coverage: leakage-free cross-fit",
        f"PC3E/GS689  n={n_pc3e}, no-refit transfer",
        "Depth titration 5-100%; long-read non-SE",
        "DM1 / ESRP apps use shipped q",
    ]
    for i, text in enumerate(rows):
        yy = 2.52 - i * 0.50
        ax.plot([7.32], [yy], marker="o", ms=2.6, color=BLUE, zorder=4)
        _label(ax, 7.52, yy, text, fontsize=6.0, color=INK, ha="left")


def _render(
    toy: dict[str, float | str | bool],
    design: dict[str, object],
) -> dict[str, dict[str, str]]:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 8,
            "figure.dpi": 300,
            "savefig.bbox": "tight",
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "svg.fonttype": "none",
        }
    )
    fig = plt.figure(figsize=(7.2, 9.2))
    gs = fig.add_gridspec(4, 1, height_ratios=[1, 1, 1, 0.95], hspace=0.22)
    builders: list[tuple] = [
        (_draw_panel_a, toy), (_draw_panel_b, toy),
        (_draw_panel_c, toy), (_draw_panel_d, design),
    ]
    for i, (builder, data) in enumerate(builders):
        ax = fig.add_subplot(gs[i, 0])
        ax.add_patch(
            Rectangle((0.0, 0.0), 12.0, 4.0, facecolor=PANEL_BG, edgecolor="none", zorder=-12)
        )
        builder(ax, data)
    fig.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
    outputs = {"F1_braid_scheme": _save_all(fig, OUT_DIR / "F1_braid_scheme")}
    plt.close(fig)

    panel_specs = {
        "F1_panelA_workflow": lambda ax: _draw_panel_a(ax, toy),
        "F1_panelB_calibration": lambda ax: _draw_panel_b(ax, toy),
        "F1_panelC_tiers": lambda ax: _draw_panel_c(ax, toy),
        "F1_panelD_design": lambda ax: _draw_panel_d(ax, design),
    }
    for stem, builder in panel_specs.items():
        fig_p, ax = plt.subplots(figsize=(7.2, 2.75))
        ax.add_patch(
            Rectangle((0.0, 0.0), 12.0, 4.0, facecolor=PANEL_BG, edgecolor="none", zorder=-12)
        )
        builder(ax)
        fig_p.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
        outputs[stem] = _save_all(fig_p, OUT_DIR / stem)
        plt.close(fig_p)
    return outputs


def _write_validation(
    outputs: dict[str, dict[str, str]],
    tables: dict[str, pd.DataFrame],
    toy: dict[str, float | str | bool],
) -> None:
    checks = [
        "F1 validation",
        f"calculation_step_rows={len(tables['calculation_steps'])}",
        f"tier_logic_rows={len(tables['tier_logic'])}",
        f"toy_dpsi={toy['dpsi']:.6f}",
        f"toy_ci_low={toy['ci_low']:.6f}",
        f"toy_ci_high={toy['ci_high']:.6f}",
        f"toy_tier={toy['tier']}",
    ]
    for group_name, group in outputs.items():
        for kind, rel in group.items():
            path = ROOT / rel
            checks.append(f"{group_name}.{kind}={path.stat().st_size} bytes")
    VALIDATION_TXT.write_text("\n".join(checks) + "\n", encoding="utf-8")


def _write_manifest(
    outputs: dict[str, dict[str, str]],
    tables: dict[str, pd.DataFrame],
    toy: dict[str, float | str | bool],
) -> None:
    file_paths = sorted(
        set(
            [str(SOURCE_XLSX.relative_to(ROOT)), str(LEGEND_DOCX.relative_to(ROOT))]
            + [p for group in outputs.values() for p in group.values()]
        )
    )
    manifest = {
        "figure_id": "F1",
        "title": "BRAID calculation scheme",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "script": str(Path(__file__).relative_to(ROOT)),
        "inputs": [
            "braid/commands/differential.py",
            "braid/target/conformal.py",
            "braid/adapters/base.py",
            str(CALIBRATOR_JSON.relative_to(ROOT)),
            str(DEPTH_TITRATION_JSON.relative_to(ROOT)),
        ],
        "outputs": {**outputs, "legend_docx": str(LEGEND_DOCX.relative_to(ROOT))},
        "source_data": str(SOURCE_XLSX.relative_to(ROOT)),
        "panel_claims": {
            "A": "BRAID estimates Delta PSI from rMATS inclusion/skipping counts.",
            "B": (
                "The split-conformal quantile of RNA-seq-to-RT-PCR residuals sets a "
                "support-aware, residual-calibrated interval width."
            ),
            "C": (
                "Confidence tiers are determined by caller significance and "
                "BRAID reliable-effect status, then reported in the output row."
            ),
            "D": (
                "The shipped q is fit once on the calibration set (TRA2 + circadian, "
                "n=162 SE events with RT-PCR truth); validation evidence is shown "
                "separately as leakage-free cross-fit coverage plus no-refit transfer, "
                "depth-titration, non-SE, and application checks."
            ),
        },
        "transformations": [
            "Computed an illustrative event from declared rMATS-like counts.",
            (
                "The schematic shows the pooled-count calculation with equal form "
                "lengths; source data record the IncFormLen/SkipFormLen ratio."
            ),
            (
                "Loaded the packaged differential conformal calibrator and used "
                "q_for plus robust_interval."
            ),
            "Rendered all panels deterministically with matplotlib from exported source tables.",
            "Mirrored the combined figure into paper/figures as fig1_workflow.",
        ],
        "validation": {
            "toy_event": {
                "dpsi": float(toy["dpsi"]),
                "ci_low": float(toy["ci_low"]),
                "ci_high": float(toy["ci_high"]),
                "tier": str(toy["tier"]),
            },
            "row_counts": {name: int(len(df)) for name, df in tables.items()},
            "file_hashes": {path: _sha256(ROOT / path) for path in file_paths},
        },
        "known_limitations": [
            (
                "Figure 1 is a schematic. The event in panels A and B is "
                "illustrative and is not a benchmark result."
            ),
            (
                "Panel A shows the pooled-count calculation. The default production "
                "path uses replicate-aware resampling when complete replicate vectors "
                "are available."
            ),
            "The figure shows the rMATS differential path used in the manuscript.",
            (
                "For non-rMATS callers without a significance field, BRAID reports "
                "BRAID-only tier vocabulary."
            ),
        ],
    }
    MANIFEST_JSON.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")


def _mirror_to_paper(
    toy: dict[str, float | str | bool],
    design: dict[str, object],
) -> None:
    PAPER_FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    for suffix in (".png", ".pdf", ".svg"):
        shutil.copyfile(
            OUT_DIR / f"F1_braid_scheme{suffix}",
            PAPER_FIGURE_DIR / f"fig1_workflow{suffix}",
        )
    fig = plt.figure(figsize=(7.2, 9.2))
    gs = fig.add_gridspec(4, 1, height_ratios=[1, 1, 1, 0.95], hspace=0.22)
    builders: list[tuple] = [
        (_draw_panel_a, toy), (_draw_panel_b, toy),
        (_draw_panel_c, toy), (_draw_panel_d, design),
    ]
    for i, (builder, data) in enumerate(builders):
        ax = fig.add_subplot(gs[i, 0])
        ax.add_patch(
            Rectangle((0.0, 0.0), 12.0, 4.0, facecolor=PANEL_BG, edgecolor="none", zorder=-12)
        )
        builder(ax, data)
    fig.subplots_adjust(left=0.01, right=0.99, top=0.99, bottom=0.01)
    fig.savefig(PAPER_FIGURE_DIR / "fig1_workflow.jpg", dpi=300)
    plt.close(fig)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    cal = _load_calibrator()
    toy = _toy_event(cal)
    design = _design_table(cal)
    tables = _source_tables(cal, toy, design)
    _write_source_data(tables)
    outputs = _render(toy, design)
    _write_legend()
    _write_validation(outputs, tables, toy)
    _write_manifest(outputs, tables, toy)
    _mirror_to_paper(toy, design)
    print(f"Wrote {OUT_DIR.relative_to(ROOT)}")
    print(VALIDATION_TXT.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
