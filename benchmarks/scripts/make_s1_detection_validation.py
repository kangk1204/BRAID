#!/usr/bin/env python3
# ruff: noqa: I001
"""Build Supplementary Figure S1 for TRA2 RT-PCR detection validation.

The figure uses the only benchmark in the current manuscript package with both
RT-PCR-positive and RT-PCR-negative event panels: GSE59335 TRA2A/B knockdown.
It converts existing head-to-head detection records into confusion matrices,
MCC values, a source-data workbook, a manifest, and a concise DOCX legend.
"""

from __future__ import annotations

import hashlib
import json
import math
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
HEADTOHEAD_DIR = ROOT / "benchmarks" / "headtohead"
if str(HEADTOHEAD_DIR) not in sys.path:
    sys.path.insert(0, str(HEADTOHEAD_DIR))
sys.path.insert(0, str(ROOT / "benchmarks"))

import detection_analysis as DA  # noqa: E402
import figstyle  # noqa: E402

figstyle.apply()


OUT_DIR = ROOT / "outputs" / "figures" / "manuscript" / "S1"
PAPER_FIGURE_DIR = ROOT / "paper" / "figures"
SOURCE_XLSX = OUT_DIR / "s1_source_data.xlsx"
MANIFEST_JSON = OUT_DIR / "s1_manifest.json"
VALIDATION_TXT = OUT_DIR / "s1_validation.txt"
LEGEND_DOCX = OUT_DIR / "s1_legend.docx"

RULES: list[tuple[str, str, Callable[[dict[str, object]], bool], str]] = [
    (
        "rmats_fdr",
        "rMATS FDR < 0.05",
        lambda r: np.isfinite(float(r["rmats_fdr"])) and float(r["rmats_fdr"]) < 0.05,
        "#CC79A7",
    ),
    (
        "posterior_effect",
        "Posterior effect",
        lambda r: float(r["prob_large"]) >= 0.5,
        "#009E73",
    ),
    (
        "braid_supported",
        "BRAID effect-supported",
        lambda r: (
            np.isfinite(float(r["rmats_fdr"]))
            and float(r["rmats_fdr"]) < 0.05
            and float(r["prob_large"]) >= 0.5
        ),
        "#0072B2",
    ),
    (
        "calibrated_ci_excludes_zero",
        "Calibrated CI excludes 0",
        lambda r: bool(r["conf_excl0"]),
        "#56B4E9",
    ),
    (
        "jeffreys_ci_excludes_zero",
        "Jeffreys CI excludes 0",
        lambda r: bool(r["jeff_excl0"]),
        "#E69F00",
    ),
]


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _save_all(fig: plt.Figure, stem: Path) -> dict[str, str]:
    outputs: dict[str, str] = {}
    for suffix in (".png", ".pdf", ".svg"):
        path = stem.with_suffix(suffix)
        if suffix == ".png":
            fig.savefig(path, dpi=300)
        else:
            fig.savefig(path)
        outputs[suffix.lstrip(".")] = str(path.relative_to(ROOT))
    return outputs


def _wilson(k: int, n: int, z: float = 1.959963984540054) -> tuple[float, float, float]:
    if n == 0:
        return float("nan"), float("nan"), float("nan")
    p = k / n
    denom = 1 + z * z / n
    center = (p + z * z / (2 * n)) / denom
    half = z * math.sqrt((p * (1 - p) + z * z / (4 * n)) / n) / denom
    return p, max(0.0, center - half), min(1.0, center + half)


def _mcc(tp: int, fn: int, fp: int, tn: int) -> float:
    denom = math.sqrt((tp + fp) * (tp + fn) * (tn + fp) * (tn + fn))
    return (tp * tn - fp * fn) / denom if denom else float("nan")


def _build_tables() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    rows = DA.build()
    labels = np.array([int(r["label"]) for r in rows], dtype=int)
    n_positive = int(np.sum(labels == 1))
    n_negative = int(np.sum(labels == 0))

    summary_rows: list[dict[str, object]] = []
    rate_rows: list[dict[str, object]] = []
    confusion_rows: list[dict[str, object]] = []
    call_rows: list[dict[str, object]] = []

    for rule_id, label, predicate, color in RULES:
        calls = np.array([bool(predicate(r)) for r in rows], dtype=bool)
        tp = int(np.sum(calls & (labels == 1)))
        fn = int(np.sum(~calls & (labels == 1)))
        fp = int(np.sum(calls & (labels == 0)))
        tn = int(np.sum(~calls & (labels == 0)))
        sensitivity, sensitivity_low, sensitivity_high = _wilson(tp, n_positive)
        fpr, fpr_low, fpr_high = _wilson(fp, n_negative)
        mcc = _mcc(tp, fn, fp, tn)

        summary_rows.append(
            {
                "rule_id": rule_id,
                "rule_label": label,
                "color": color,
                "n_positive": n_positive,
                "n_negative": n_negative,
                "tp": tp,
                "fn": fn,
                "fp": fp,
                "tn": tn,
                "sensitivity": sensitivity,
                "sensitivity_wilson_low": sensitivity_low,
                "sensitivity_wilson_high": sensitivity_high,
                "false_positive_rate": fpr,
                "false_positive_rate_wilson_low": fpr_low,
                "false_positive_rate_wilson_high": fpr_high,
                "specificity": tn / n_negative if n_negative else float("nan"),
                "mcc": mcc,
            }
        )
        for metric, numerator, denominator, estimate, low, high in (
            ("Sensitivity", tp, n_positive, sensitivity, sensitivity_low, sensitivity_high),
            ("False-positive rate", fp, n_negative, fpr, fpr_low, fpr_high),
        ):
            rate_rows.append(
                {
                    "rule_id": rule_id,
                    "rule_label": label,
                    "metric": metric,
                    "numerator": numerator,
                    "denominator": denominator,
                    "estimate": estimate,
                    "wilson_low": low,
                    "wilson_high": high,
                }
            )
        for truth_label, call_label, value in (
            ("RT-PCR positive", "Called positive", tp),
            ("RT-PCR positive", "Not called", fn),
            ("RT-PCR negative", "Called positive", fp),
            ("RT-PCR negative", "Not called", tn),
        ):
            confusion_rows.append(
                {
                    "rule_id": rule_id,
                    "rule_label": label,
                    "truth": truth_label,
                    "call": call_label,
                    "count": value,
                }
            )
        for i, (r, call) in enumerate(zip(rows, calls, strict=True), start=1):
            call_rows.append(
                {
                    "row_index": i,
                    "rule_id": rule_id,
                    "rule_label": label,
                    "rtpcr_label": "positive" if int(r["label"]) == 1 else "negative",
                    "called_positive": bool(call),
                    "rmats_fdr": float(r["rmats_fdr"]),
                    "posterior_prob_abs_dpsi_ge_0_1": float(r["prob_large"]),
                    "calibrated_ci_excludes_zero": bool(r["conf_excl0"]),
                    "jeffreys_ci_excludes_zero": bool(r["jeff_excl0"]),
                    "replicate_z": float(r["zsc"]),
                    "dpsi_mean": float(r["dpsi"]),
                }
            )

    metadata = pd.DataFrame(
        [
            {"key": "figure", "value": "Supplementary Figure S1"},
            {"key": "generated_by", "value": str(Path(__file__).relative_to(ROOT))},
            {"key": "source_script", "value": "benchmarks/headtohead/detection_analysis.py"},
            {"key": "dataset", "value": "GSE59335 TRA2A/B knockdown"},
            {"key": "matched_targets", "value": str(len(rows))},
            {"key": "rtpcr_positive_targets", "value": str(n_positive)},
            {"key": "rtpcr_negative_targets", "value": str(n_negative)},
            {
                "key": "denominator",
                "value": "matched RT-PCR targets with usable rMATS counts",
            },
        ]
    )
    return (
        pd.DataFrame(summary_rows),
        pd.DataFrame(rate_rows),
        pd.DataFrame(confusion_rows),
        pd.DataFrame(call_rows),
        metadata,
    )


def _plot_panel_a(ax: plt.Axes, summary: pd.DataFrame) -> None:
    y = np.arange(len(summary))
    ax.errorbar(
        summary["sensitivity"],
        y + 0.11,
        xerr=[
            summary["sensitivity"] - summary["sensitivity_wilson_low"],
            summary["sensitivity_wilson_high"] - summary["sensitivity"],
        ],
        fmt="o",
        color="#0072B2",
        capsize=2.5,
        lw=1.2,
        label="Sensitivity",
    )
    ax.errorbar(
        summary["false_positive_rate"],
        y - 0.11,
        xerr=[
            summary["false_positive_rate"] - summary["false_positive_rate_wilson_low"],
            summary["false_positive_rate_wilson_high"] - summary["false_positive_rate"],
        ],
        fmt="s",
        color="#D55E00",
        capsize=2.5,
        lw=1.2,
        label="False-positive rate",
    )
    ax.set_yticks(y, summary["rule_label"])
    ax.set_xlim(-0.03, 1.03)
    ax.set_xlabel("Rate with Wilson 95% CI")
    ax.set_title("A  Detection rates")
    ax.invert_yaxis()
    ax.grid(axis="x", color="#dddddd", linewidth=0.6)
    ax.legend(
        frameon=False,
        loc="upper center",
        bbox_to_anchor=(0.5, -0.16),
        ncol=2,
        borderaxespad=0.0,
        fontsize=7,
    )


def _plot_panel_b(ax: plt.Axes, summary: pd.DataFrame) -> None:
    x = np.arange(len(summary))
    colors = list(summary["color"])
    values = summary["mcc"].astype(float).to_numpy()
    ax.bar(x, values, color=colors, width=0.68)
    ax.axhline(0, color="black", lw=0.8)
    ax.set_xticks(x, summary["rule_label"], rotation=28, ha="right")
    ax.set_ylim(0, max(values) * 1.22)
    ax.set_ylabel("Matthews correlation coefficient")
    ax.set_title("B  Balanced detection accuracy")
    for i, value in enumerate(values):
        ax.text(i, value + 0.015, f"{value:.2f}", ha="center", va="bottom", fontsize=7)


def _confusion_matrix(summary: pd.DataFrame, rule_id: str) -> np.ndarray:
    row = summary.loc[summary["rule_id"] == rule_id].iloc[0]
    return np.array([[row["tp"], row["fn"]], [row["fp"], row["tn"]]], dtype=float)


def _plot_confusion(ax: plt.Axes, summary: pd.DataFrame, rule_id: str, title: str) -> None:
    matrix = _confusion_matrix(summary, rule_id)
    vmax = float(matrix.max())
    ax.imshow(matrix, cmap="Blues", vmin=0, vmax=vmax)
    ax.set_xticks([0, 1], ["Called\npositive", "Not\ncalled"])
    ax.set_yticks([0, 1], ["RT-PCR\npositive", "RT-PCR\nnegative"])
    ax.set_title(title)
    labels = np.array([["TP", "FN"], ["FP", "TN"]])
    for i in range(2):
        for j in range(2):
            value = int(matrix[i, j])
            text_color = "white" if matrix[i, j] > vmax * 0.55 else "black"
            ax.text(
                j,
                i,
                f"{labels[i, j]}\n{value}",
                ha="center",
                va="center",
                color=text_color,
                fontsize=8,
                weight="bold",
            )
    for spine in ax.spines.values():
        spine.set_visible(False)


def _render(summary: pd.DataFrame) -> dict[str, dict[str, str]]:
    plt.rcParams.update(
        {
            "font.size": 8,
            "axes.titlesize": 9,
            "axes.labelsize": 8,
            "xtick.labelsize": 7,
            "ytick.labelsize": 7,
            "legend.fontsize": 7,
            "font.family": "DejaVu Sans",
        }
    )
    fig = plt.figure(figsize=(7.2, 6.8), constrained_layout=True)
    gs = fig.add_gridspec(2, 2, height_ratios=[1.1, 0.9])
    ax_a = fig.add_subplot(gs[0, 0])
    ax_b = fig.add_subplot(gs[0, 1])
    ax_c = fig.add_subplot(gs[1, 0])
    ax_d = fig.add_subplot(gs[1, 1])
    _plot_panel_a(ax_a, summary)
    _plot_panel_b(ax_b, summary)
    _plot_confusion(ax_c, summary, "rmats_fdr", "C  rMATS FDR < 0.05")
    _plot_confusion(ax_d, summary, "braid_supported", "D  BRAID effect-supported")
    outputs = {"S1_detection_validation": _save_all(fig, OUT_DIR / "S1_detection_validation")}
    plt.close(fig)

    # Standalone panels for journal layout flexibility.
    panel_builders = {
        "S1_panelA_detection_rates": lambda ax: _plot_panel_a(ax, summary),
        "S1_panelB_mcc": lambda ax: _plot_panel_b(ax, summary),
        "S1_panelC_rmats_confusion": lambda ax: _plot_confusion(
            ax, summary, "rmats_fdr", "C  rMATS FDR < 0.05"
        ),
        "S1_panelD_braid_confusion": lambda ax: _plot_confusion(
            ax, summary, "braid_supported", "D  BRAID effect-supported"
        ),
    }
    for name, builder in panel_builders.items():
        fig_panel, ax_panel = plt.subplots(figsize=(3.5, 3.0), constrained_layout=True)
        builder(ax_panel)
        outputs[name] = _save_all(fig_panel, OUT_DIR / name)
        plt.close(fig_panel)
    return outputs


def _write_source_data(
    summary: pd.DataFrame,
    rates: pd.DataFrame,
    confusion: pd.DataFrame,
    calls: pd.DataFrame,
    metadata: pd.DataFrame,
) -> None:
    with pd.ExcelWriter(SOURCE_XLSX, engine="openpyxl") as writer:
        rates.to_excel(writer, index=False, sheet_name="panel_A_rates")
        summary.to_excel(writer, index=False, sheet_name="panel_B_mcc")
        confusion.to_excel(writer, index=False, sheet_name="panel_CD_confusion")
        calls.to_excel(writer, index=False, sheet_name="per_event_calls")
        metadata.to_excel(writer, index=False, sheet_name="metadata")


def _write_legend(summary: pd.DataFrame) -> None:
    braid = summary.loc[summary["rule_id"] == "braid_supported"].iloc[0]
    rmats = summary.loc[summary["rule_id"] == "rmats_fdr"].iloc[0]
    doc = Document()
    doc.add_paragraph(
        "Supplementary Figure S1. Auxiliary TRA2 positive/negative detection check."
    )
    doc.add_paragraph(
        "A, Sensitivity and false-positive rate across detection rules on 76 "
        "RT-PCR-positive and 36 RT-PCR-negative matched GSE59335 cassette-exon "
        "targets; horizontal bars show Wilson 95% confidence intervals. B, MCC "
        "is shown only as a secondary summary metric for the same TRA2-only "
        "rules. C,D, Confusion matrices for rMATS FDR < 0.05 and BRAID effect-supported "
        "calls. BRAID effect-supported "
        "combines rMATS FDR < 0.05 with posterior P(|Delta PSI| >= 0.1) >= 0.5 "
        f"and reduced false positives from {int(rmats['fp'])} to {int(braid['fp'])} "
        f"with MCC changing from {float(rmats['mcc']):.2f} to "
        f"{float(braid['mcc']):.2f}."
    )
    doc.save(LEGEND_DOCX)


def _write_manifest(
    outputs: dict[str, dict[str, str]],
    summary: pd.DataFrame,
    rates: pd.DataFrame,
    confusion: pd.DataFrame,
    calls: pd.DataFrame,
) -> None:
    manifest_files = {
        "source_data": str(SOURCE_XLSX.relative_to(ROOT)),
        "legend_docx": str(LEGEND_DOCX.relative_to(ROOT)),
        "validation": str(VALIDATION_TXT.relative_to(ROOT)),
    }
    for group in outputs.values():
        manifest_files.update({path: path for path in group.values()})

    braid = summary.loc[summary["rule_id"] == "braid_supported"].iloc[0]
    rmats = summary.loc[summary["rule_id"] == "rmats_fdr"].iloc[0]
    manifest = {
        "figure_id": "S1",
        "title": "Auxiliary TRA2 positive/negative detection check",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "script": str(Path(__file__).relative_to(ROOT)),
        "inputs": [
            "benchmarks/headtohead/detection_analysis.py",
            "benchmarks/headtohead/head_to_head_coverage.py",
            "data/public_benchmarks/GSE59335/rmats/SE.MATS.JC.txt",
            "data/public_benchmarks/GSE59335/targets/validated_events.tsv",
            "data/public_benchmarks/GSE59335/targets/failed_events.tsv",
        ],
        "outputs": outputs,
        "source_data": str(SOURCE_XLSX.relative_to(ROOT)),
        "legend_docx": str(LEGEND_DOCX.relative_to(ROOT)),
        "panel_claims": {
            "A": (
                "Detection sensitivity and false-positive rate were evaluated only "
                "on the matched TRA2 positive/negative RT-PCR target set."
            ),
            "B": (
                "MCC is a secondary summary metric for the TRA2-only matched "
                "positive/negative check."
            ),
            "C": "rMATS FDR < 0.05 calls 8 of 36 RT-PCR-negative targets positive.",
            "D": "BRAID effect-supported calls 2 of 36 RT-PCR-negative targets positive.",
        },
        "transformations": [
            "Matched RT-PCR cassette-exon targets to GSE59335 rMATS skipped-exon events.",
            (
                "Computed detection calls from rMATS FDR, posterior large-effect "
                "probability, Jeffreys intervals, and calibrated intervals."
            ),
            (
                "Computed sensitivity, false-positive rate, Wilson 95% intervals, "
                "confusion matrices, and MCC as a secondary summary metric."
            ),
        ],
        "validation": {
            "matched_targets": int(calls["row_index"].max()),
            "n_positive": int(summary["n_positive"].iloc[0]),
            "n_negative": int(summary["n_negative"].iloc[0]),
            "rate_rows": int(len(rates)),
            "summary_rows": int(len(summary)),
            "confusion_rows": int(len(confusion)),
            "per_event_call_rows": int(len(calls)),
            "rmats_mcc": float(rmats["mcc"]),
            "braid_supported_mcc": float(braid["mcc"]),
            "rmats_false_positives": int(rmats["fp"]),
            "braid_supported_false_positives": int(braid["fp"]),
            "file_hashes": {
                path: _sha256(ROOT / path)
                for path in sorted(
                    set(
                        [str(SOURCE_XLSX.relative_to(ROOT)), str(LEGEND_DOCX.relative_to(ROOT))]
                        + [p for group in outputs.values() for p in group.values()]
                    )
                )
            },
        },
        "known_limitations": [
            (
                "MCC is a secondary metric reported only for GSE59335 because "
                "the other manuscript RT-PCR benchmarks are positive-only target "
                "panels."
            ),
            (
                "The denominator is the matched target set with usable rMATS counts, "
                "not every row in the original RT-PCR tables."
            ),
            (
                "Posterior-effect-only calls are shown as an ablation, not as the "
                "production BRAID tier."
            ),
        ],
    }
    MANIFEST_JSON.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")


def _write_validation(summary: pd.DataFrame, outputs: dict[str, dict[str, str]]) -> None:
    braid = summary.loc[summary["rule_id"] == "braid_supported"].iloc[0]
    rmats = summary.loc[summary["rule_id"] == "rmats_fdr"].iloc[0]
    lines = [
        "S1 validation",
        f"n_positive={int(summary['n_positive'].iloc[0])}",
        f"n_negative={int(summary['n_negative'].iloc[0])}",
        f"rmats_tp_fn_fp_tn={int(rmats['tp'])},{int(rmats['fn'])},{int(rmats['fp'])},{int(rmats['tn'])}",
        f"rmats_mcc={float(rmats['mcc']):.6f}",
        (
            "braid_supported_tp_fn_fp_tn="
            f"{int(braid['tp'])},{int(braid['fn'])},{int(braid['fp'])},{int(braid['tn'])}"
        ),
        f"braid_supported_mcc={float(braid['mcc']):.6f}",
    ]
    for group_name, group in outputs.items():
        for kind, rel in group.items():
            path = ROOT / rel
            lines.append(f"{group_name}.{kind}={path.stat().st_size} bytes")
    VALIDATION_TXT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    PAPER_FIGURE_DIR.mkdir(parents=True, exist_ok=True)

    summary, rates, confusion, calls, metadata = _build_tables()
    _write_source_data(summary, rates, confusion, calls, metadata)
    outputs = _render(summary)
    _write_legend(summary)
    _write_validation(summary, outputs)
    _write_manifest(outputs, summary, rates, confusion, calls)

    for suffix in (".png", ".pdf", ".svg"):
        shutil.copyfile(
            OUT_DIR / f"S1_detection_validation{suffix}",
            PAPER_FIGURE_DIR / f"supp_fig1_detection_validation{suffix}",
        )

    print(f"Wrote {OUT_DIR.relative_to(ROOT)}")
    print(VALIDATION_TXT.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
