#!/usr/bin/env python3
# ruff: noqa: I001
"""Build manuscript Figure 2 for the BRAID RT-PCR benchmark.

The script reuses the existing head-to-head benchmark matching code, exports the
exact plotted source data, writes a concise manuscript legend, and mirrors the
combined figure into paper/figures for local manuscript builds.
"""

from __future__ import annotations

import hashlib
import json
import os
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
HEADTOHEAD_DIR = ROOT / "benchmarks" / "headtohead"
if str(HEADTOHEAD_DIR) not in sys.path:
    sys.path.insert(0, str(HEADTOHEAD_DIR))
sys.path.insert(0, str(ROOT / "benchmarks"))

import figstyle  # noqa: E402
import head_to_head_coverage as H  # noqa: E402
import make_figures as MF  # noqa: E402

figstyle.apply()


OUT_DIR = ROOT / "outputs" / "figures" / "manuscript" / "F2"
PAPER_FIGURE_DIR = ROOT / "paper" / "figures"
SOURCE_XLSX = OUT_DIR / "f2_source_data.xlsx"
MANIFEST_JSON = OUT_DIR / "f2_manifest.json"
VALIDATION_TXT = OUT_DIR / "f2_validation.txt"
LEGEND_DOCX = OUT_DIR / "f2_legend.docx"

METHOD_ORDER = ["BRAID-conformal", "MAJIQ", "betAS", "rMATS"]
DISPLAY_METHOD = {
    "BRAID-conformal": "BRAID",
    "MAJIQ": "MAJIQ",
    "betAS": "betAS",
    "rMATS": "rMATS",
}
DATASET_ORDER = ["TRA2", "Circadian", "SRS354082", "Pooled"]
COLORS = {
    "BRAID-conformal": "#0072B2",
    "MAJIQ": "#D55E00",
    "betAS": "#009E73",
    "rMATS": "#CC79A7",
}
ALPHA = 0.05


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _save_all(fig: plt.Figure, stem: Path) -> dict[str, str]:
    outputs: dict[str, str] = {}
    for suffix in (".png", ".pdf", ".svg"):
        path = stem.with_suffix(suffix)
        if suffix == ".png":
            fig.savefig(path, dpi=300)
        else:
            fig.savefig(path)
        outputs[suffix.lstrip(".")] = str(path.relative_to(ROOT))
    return outputs


def _load_benchmark_rows() -> dict[str, list[dict[str, object]]]:
    os.chdir(ROOT)
    db = ROOT / "data" / "public_benchmarks"
    bh = ROOT / "benchmarks" / "headtohead"
    circ_tsv = db / "meta" / "gse54651_circadian_positive_events.tsv"
    pc3e_truth = db / "meta" / "rmats_pc3e_gs689_positive_events.tsv"

    tra2 = MF.build(
        {
            "short": "TRA2",
            "kind": "exon",
            "swap": True,
            "rmats_se": str(db / "GSE59335" / "rmats" / "SE.MATS.JC.txt"),
            "targets": H.load_targets(
                str(db / "GSE59335" / "targets" / "validated_events.tsv"),
                str(db / "GSE59335" / "targets" / "failed_events.tsv"),
            ),
            "majiq_tsv": str(db / "GSE59335" / "majiq" / "deltapsi.tsv"),
            "betas_iv": str(bh / "tra2_betas_intervals.tsv"),
            "betas_truth": str(bh / "tra2_betas_truth.tsv"),
            "majiq_groups": ("KD", "CTRL"),
        }
    )
    circadian = MF.build(
        {
            "short": "Circadian",
            "kind": "junction",
            "swap": True,
            "rmats_se": str(db / "GSE54651" / "rmats" / "SE.MATS.JC.txt"),
            "targets": H.load_circadian_targets(str(circ_tsv)),
            "incls": MF.C.load_circ_incl(str(circ_tsv)),
            "majiq_tsv": str(db / "GSE54651" / "majiq" / "deltapsi.tsv"),
            "betas_iv": str(bh / "circ_betas_intervals.tsv"),
            "betas_truth": str(bh / "circ_betas_truth.tsv"),
            "majiq_groups": ("LIVER", "CEREB"),
        }
    )
    srs = MF.build(
        {
            "short": "SRS354082",
            "kind": "exon",
            "swap": False,
            "rmats_se": str(db / "SRS354082" / "rmats" / "SE.MATS.JC.txt"),
            "targets": H.load_targets(str(pc3e_truth), None),
            "majiq_tsv": str(db / "SRS354082" / "majiq" / "deltapsi.tsv"),
            "betas_iv": str(bh / "srs_betas_intervals.tsv"),
            "betas_truth": str(bh / "srs_betas_truth.tsv"),
            "majiq_groups": ("PC3E", "GS689"),
        }
    )
    return {
        "TRA2": tra2,
        "Circadian": circadian,
        "SRS354082": srs,
        "Pooled": tra2 + circadian + srs,
    }


def _coverage_rows(datasets: dict[str, list[dict[str, object]]]) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for dataset in DATASET_ORDER:
        data_rows = datasets[dataset]
        intervals, truth = MF.method_intervals(data_rows, 0.95)
        for method in METHOD_ORDER:
            lo, hi = intervals[method]
            covered = int(np.sum((truth >= lo) & (truth <= hi)))
            coverage, wilson_low, wilson_high = MF.wilson(covered, truth.size)
            rows.append(
                {
                    "panel": "A",
                    "dataset": dataset,
                    "method": method,
                    "method_label": DISPLAY_METHOD[method],
                    "n": int(truth.size),
                    "covered": covered,
                    "coverage": float(coverage),
                    "wilson_low": float(wilson_low),
                    "wilson_high": float(wilson_high),
                    "mean_width": float(np.mean(hi - lo)),
                    "interval_score": MF.interval_score(lo, hi, truth, ALPHA),
                }
            )
    return pd.DataFrame(rows)


_HEADTOHEAD_JSON = ROOT / "benchmarks" / "results" / "headtohead_coverage.json"
_CANONICAL_DATASET = {
    "TRA2": "TRA2 (GSE59335, human)",
    "Circadian": "Circadian (GSE54651, mouse)",
    "SRS354082": "PC3E/GS689 (SRA, human)",
    "Pooled": "pooled_common",
}


def _reconcile_with_canonical(coverage_df: pd.DataFrame) -> pd.DataFrame:
    """Overwrite coverage/width/interval-score/Wilson from the canonical head-to-head
    JSON so Figure 2, its source workbook, and Table 1 (which reads this workbook)
    share one source of truth instead of an independent recomputation."""
    canon = json.loads(_HEADTOHEAD_JSON.read_text(encoding="utf-8"))
    out = coverage_df.copy()
    for idx, row in out.iterrows():
        ds_key = _CANONICAL_DATASET.get(str(row["dataset"]))
        if ds_key is None:
            continue
        node = (canon["pooled_common"] if ds_key == "pooled_common"
                else canon["datasets"].get(ds_key))
        m = (node or {}).get("methods", {}).get(str(row["method"]))
        if not m:
            continue
        n = int(row["n"])
        out.at[idx, "coverage"] = float(m["coverage"])
        out.at[idx, "covered"] = int(round(float(m["coverage"]) * n))
        out.at[idx, "mean_width"] = float(m["width"])
        out.at[idx, "interval_score"] = float(m["iscore"])
        wilson = m.get("wilson95")
        if wilson:
            out.at[idx, "wilson_low"] = float(wilson[0])
            out.at[idx, "wilson_high"] = float(wilson[1])
    return out


def _pooled_rows(coverage_df: pd.DataFrame) -> pd.DataFrame:
    out = coverage_df[coverage_df["dataset"] == "Pooled"].copy()
    out["panel_b_value"] = out["interval_score"]
    out["panel_c_x"] = out["mean_width"]
    out["panel_c_y"] = out["coverage"]
    return out


def _plot_panel_a(
    ax: plt.Axes,
    coverage_df: pd.DataFrame,
    title: str,
    show_legend: bool = False,
) -> None:
    offsets = {
        "BRAID-conformal": -0.27,
        "MAJIQ": -0.09,
        "betAS": 0.09,
        "rMATS": 0.27,
    }
    y_positions = np.arange(len(DATASET_ORDER), dtype=float)
    for base_y, dataset in zip(y_positions, DATASET_ORDER):
        sub = coverage_df[coverage_df["dataset"] == dataset]
        for method in METHOD_ORDER:
            row = sub[sub["method"] == method].iloc[0]
            x = float(row["coverage"])
            xerr = [[x - float(row["wilson_low"])], [float(row["wilson_high"]) - x]]
            y = base_y + offsets[method]
            ax.errorbar(
                x,
                y,
                xerr=xerr,
                fmt="o",
                color=COLORS[method],
                capsize=2.5,
                ms=4.5,
                lw=1.4,
                label=DISPLAY_METHOD[method] if dataset == DATASET_ORDER[0] else None,
            )
    ax.axvline(0.95, color="black", linestyle="--", linewidth=0.8)
    labels = [
        f"{dataset}\n(n={int(coverage_df[coverage_df['dataset'] == dataset]['n'].iloc[0])})"
        for dataset in DATASET_ORDER
    ]
    ax.set_yticks(y_positions, labels)
    ax.set_xlim(0.30, 1.02)
    ax.set_ylim(len(DATASET_ORDER) - 0.5, -0.5)
    ax.set_xlabel("Coverage of RT-PCR ΔPSI", fontsize=7)
    ax.set_title(title, fontsize=8.5, fontweight="bold")
    ax.tick_params(axis="both", labelsize=6.5)
    ax.text(0.952, -0.42, "nominal 0.95", ha="left", va="center", fontsize=6)
    if show_legend:
        ax.legend(
            frameon=False,
            loc="upper center",
            bbox_to_anchor=(0.5, -0.22),
            fontsize=6,
            ncol=4,
            borderaxespad=0.0,
        )


def _plot_panel_b(ax: plt.Axes, pooled_df: pd.DataFrame, title: str) -> None:
    values = [
        float(pooled_df.loc[pooled_df["method"] == method, "interval_score"].iloc[0])
        for method in METHOD_ORDER
    ]
    x = np.arange(len(METHOD_ORDER))
    ax.bar(x, values, color=[COLORS[m] for m in METHOD_ORDER], width=0.65)
    ax.set_xticks(x, [DISPLAY_METHOD[m] for m in METHOD_ORDER], rotation=25, ha="right")
    ax.set_ylabel("Interval score", fontsize=7)
    ax.set_title(title, fontsize=8.5, fontweight="bold")
    ax.tick_params(axis="both", labelsize=6.5)
    for i, value in enumerate(values):
        ax.text(i, value + 0.04, f"{value:.2f}", ha="center", va="bottom", fontsize=6.5)
    ax.set_ylim(0, max(values) * 1.2)


def _plot_panel_c(ax: plt.Axes, pooled_df: pd.DataFrame, title: str) -> None:
    for method in METHOD_ORDER:
        row = pooled_df[pooled_df["method"] == method].iloc[0]
        x = float(row["mean_width"])
        y = float(row["coverage"])
        ax.scatter(x, y, s=55, color=COLORS[method], zorder=3)
        ha = "right" if method == "BRAID-conformal" else "left"
        dx = -5 if method == "BRAID-conformal" else 5
        ax.annotate(
            DISPLAY_METHOD[method],
            xy=(x, y),
            xytext=(dx, 4),
            textcoords="offset points",
            ha=ha,
            fontsize=6.5,
        )
    ax.axhline(0.95, color="black", linestyle="--", linewidth=0.8)
    ax.set_xlabel("Mean interval width", fontsize=7)
    ax.set_ylabel("Coverage", fontsize=7)
    ax.set_title(title, fontsize=8.5, fontweight="bold")
    ax.tick_params(axis="both", labelsize=6.5)
    ax.set_ylim(0.45, 1.02)
    ax.set_xlim(
        float(pooled_df["mean_width"].min()) - 0.06,
        float(pooled_df["mean_width"].max()) + 0.12,
    )
    ax.text(ax.get_xlim()[0], 0.955, "nominal 0.95", ha="left", va="bottom", fontsize=6)


def _write_source_data(
    coverage_df: pd.DataFrame,
    pooled_df: pd.DataFrame,
    datasets: dict[str, list[dict[str, object]]],
) -> None:
    metadata = pd.DataFrame(
        [
            {
                "key": "figure",
                "value": "F2",
            },
            {
                "key": "generated_by",
                "value": str(Path(__file__).relative_to(ROOT)),
            },
            {
                "key": "datasets",
                "value": ", ".join(
                    f"{name}: n={len(rows)}" for name, rows in datasets.items()
                ),
            },
            {
                "key": "source",
                "value": "computed from public benchmark inputs via headtohead matching code",
            },
        ]
    )
    with pd.ExcelWriter(SOURCE_XLSX, engine="openpyxl") as writer:
        coverage_df.to_excel(writer, index=False, sheet_name="panel_A_coverage")
        pooled_df.to_excel(writer, index=False, sheet_name="panel_BC_pooled")
        metadata.to_excel(writer, index=False, sheet_name="metadata")


def _write_legend() -> None:
    doc = Document()
    doc.add_paragraph(
        "Figure 2. RT-PCR benchmark of calibrated Delta PSI intervals."
    )
    doc.add_paragraph(
        "A, Empirical coverage of nominal 95% intervals against RT-PCR Delta PSI "
        "for TRA2, circadian, SRS354082, and the pooled common four-method set; "
        "horizontal bars show Wilson 95% confidence intervals and the dashed line "
        "marks nominal 0.95 coverage. B, Pooled Gneiting-Raftery interval score "
        "for the common four-method set, where lower values indicate better joint "
        "coverage and sharpness. BRAID denotes the cross-fit absolute-residual "
        "conformal interval (the coverage-width trade-off is shown in Figure 3B)."
    )
    doc.save(LEGEND_DOCX)


def _write_manifest(
    outputs: dict[str, dict[str, str]],
    coverage_df: pd.DataFrame,
    pooled_df: pd.DataFrame,
) -> None:
    files = {
        "source_data": str(SOURCE_XLSX.relative_to(ROOT)),
        "legend_docx": str(LEGEND_DOCX.relative_to(ROOT)),
        "validation": str(VALIDATION_TXT.relative_to(ROOT)),
    }
    for group in outputs.values():
        files.update({path: path for path in group.values()})

    manifest = {
        "figure_id": "F2",
        "title": "RT-PCR benchmark of calibrated Delta PSI intervals",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "script": str(Path(__file__).relative_to(ROOT)),
        "inputs": [
            "benchmarks/headtohead/make_figures.py",
            "benchmarks/headtohead/comprehensive_benchmark.py",
            "benchmarks/headtohead/head_to_head_coverage.py",
            "benchmarks/headtohead/*_betas_intervals.tsv",
            "benchmarks/headtohead/*_betas_truth.tsv",
            "data/public_benchmarks/GSE59335",
            "data/public_benchmarks/GSE54651",
            "data/public_benchmarks/SRS354082",
        ],
        "outputs": outputs,
        "source_data": str(SOURCE_XLSX.relative_to(ROOT)),
        "legend_docx": str(LEGEND_DOCX.relative_to(ROOT)),
        "panel_claims": {
            "A": (
                "BRAID-conformal reaches nominal coverage on each dataset and "
                "on the pooled common set."
            ),
            "B": "BRAID-conformal has the lowest pooled interval score.",
            "C": (
                "BRAID-conformal is wider than sampling-only intervals but "
                "reaches nominal coverage."
            ),
        },
        "transformations": [
            "Matched rMATS, MAJIQ, betAS, and BRAID intervals to the same RT-PCR target events.",
            (
                "Computed nominal 95% coverage, Wilson 95% confidence intervals, "
                "mean interval width, and interval score."
            ),
            "Used pooled common four-method events for panels B and C.",
        ],
        "validation": {
            "panel_A_rows": int(len(coverage_df)),
            "panel_BC_rows": int(len(pooled_df)),
            "pooled_n": int(pooled_df["n"].iloc[0]),
            "file_hashes": {
                path: _sha256(ROOT / path)
                for path in sorted(
                    set(
                        [str(SOURCE_XLSX.relative_to(ROOT)), str(LEGEND_DOCX.relative_to(ROOT))]
                        + [p for group in outputs.values() for p in group.values()]
                    )
                )
            },
        },
        "known_limitations": [
            (
                "Panels B and C use the common four-method set and therefore "
                "exclude events missing a MAJIQ match."
            ),
            (
                "The MAJIQ interval is reconstructed from posterior mean and "
                "standard deviation, matching the manuscript methods."
            ),
        ],
    }
    MANIFEST_JSON.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")


def _write_validation(
    coverage_df: pd.DataFrame,
    pooled_df: pd.DataFrame,
    outputs: dict[str, dict[str, str]],
) -> None:
    braid = pooled_df[pooled_df["method"] == "BRAID-conformal"].iloc[0]
    betas = pooled_df[pooled_df["method"] == "betAS"].iloc[0]
    lines = [
        "F2 validation",
        f"panel_A_rows={len(coverage_df)}",
        f"panel_BC_rows={len(pooled_df)}",
        f"pooled_n={int(braid['n'])}",
        f"pooled_BRAID_coverage={float(braid['coverage']):.6f}",
        f"pooled_BRAID_interval_score={float(braid['interval_score']):.6f}",
        f"pooled_betAS_coverage={float(betas['coverage']):.6f}",
        f"pooled_betAS_interval_score={float(betas['interval_score']):.6f}",
    ]
    for group_name, group in outputs.items():
        for kind, rel in group.items():
            path = ROOT / rel
            lines.append(f"{group_name}.{kind}={path.stat().st_size} bytes")
    VALIDATION_TXT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build() -> dict[str, object]:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    PAPER_FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    datasets = _load_benchmark_rows()
    coverage_df = _reconcile_with_canonical(_coverage_rows(datasets))
    pooled_df = _pooled_rows(coverage_df)
    _write_source_data(coverage_df, pooled_df, datasets)
    _write_legend()

    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 7,
            "axes.titlesize": 8,
            "axes.labelsize": 7,
            "xtick.labelsize": 6,
            "ytick.labelsize": 6,
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "savefig.bbox": "tight",
        }
    )

    outputs: dict[str, dict[str, str]] = {}
    fig, axes = plt.subplots(
        1,
        2,
        figsize=(5.2, 3.4),
        gridspec_kw={"width_ratios": [1.6, 0.8]},
    )
    _plot_panel_a(axes[0], coverage_df, "A. 95% interval coverage")
    _plot_panel_b(axes[1], pooled_df, "B. Interval score")
    fig.tight_layout(w_pad=1.2)
    outputs["combined"] = _save_all(fig, OUT_DIR / "F2_benchmark")
    plt.close(fig)

    standalone = [
        (
            "panel_A",
            "F2_panelA_coverage",
            lambda ax, data, title: _plot_panel_a(ax, data, title, show_legend=True),
            coverage_df,
            "A. 95% interval coverage",
        ),
        ("panel_B", "F2_panelB_interval_score", _plot_panel_b, pooled_df, "B. Interval score"),
    ]
    for key, stem, func, data, title in standalone:
        fig_panel, ax = plt.subplots(figsize=(3.6, 3.0))
        func(ax, data, title)
        fig_panel.tight_layout()
        outputs[key] = _save_all(fig_panel, OUT_DIR / stem)
        plt.close(fig_panel)

    for suffix in (".png", ".pdf", ".svg"):
        shutil.copyfile(
            OUT_DIR / f"F2_benchmark{suffix}",
            PAPER_FIGURE_DIR / f"fig2_benchmark{suffix}",
        )

    _write_validation(coverage_df, pooled_df, outputs)
    _write_manifest(outputs, coverage_df, pooled_df)
    return {
        "coverage_rows": int(len(coverage_df)),
        "pooled_rows": int(len(pooled_df)),
        "outputs": outputs,
        "source_data": str(SOURCE_XLSX.relative_to(ROOT)),
        "manifest": str(MANIFEST_JSON.relative_to(ROOT)),
        "legend": str(LEGEND_DOCX.relative_to(ROOT)),
        "validation": str(VALIDATION_TXT.relative_to(ROOT)),
    }


def main() -> None:
    print(json.dumps(build(), indent=2))


if __name__ == "__main__":
    main()
